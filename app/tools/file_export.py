"""Server `file-export` (plugin) — tool THẬT xuất kết quả ra file để user tải về.

Tách bạch flow gốc: chỉ là 1 ToolProvider trong catalog. Agent CHỈ thấy/gọi được
khi được Master gắn connector `file-export`; mỗi lượt model tự quyết có gọi không
(đúng nguyên tắc "có nhu cầu mới gọi tới").

Cơ chế giao file dùng chung với Upia: tool ghi file ra đĩa theo conversation_id rồi
trả JSON `{"artifact": true, "filename": ...}`. ChatEngine đọc base64 + emit event
`artifact` → frontend render nút tải (không cần model in link).

Định dạng hỗ trợ (đợt này): CSV (stdlib), Excel .xlsx (openpyxl), Word .docx (python-docx).
Các lib đã có sẵn trong deps (dùng cho upload). KHÔNG thêm dependency mới.
"""

import base64
import csv
import io
import json
import logging
import re
from pathlib import Path
from typing import Any

from app.llm.base import ToolDef

log = logging.getLogger(__name__)

# Artifact đặt dưới data/ (đã gitignore) — KHÔNG bẩn repo. parents[2] = repo root
# (app/tools/file_export.py → app → root).
_DATA_ROOT = Path(__file__).resolve().parents[2] / "data"
_EXPORT_ROOT = _DATA_ROOT / "export_artifacts"

# Giới hạn an toàn cho model nhỏ tự sinh nội dung (tránh cắt token giữa chừng / file khổng lồ).
_MAX_ROWS = 500
_MAX_COLS = 50

# Đuôi file hợp lệ theo từng tool (ép đúng để frontend nhận đúng MIME).
_EXT_BY_TOOL = {"export_csv": ".csv", "export_xlsx": ".xlsx", "export_docx": ".docx"}


def _safe_segment(value: str, fallback: str) -> str:
    """Chuẩn hoá 1 đoạn tên dùng trong path (conversation_id) — chỉ [a-zA-Z0-9_-]."""
    cleaned = re.sub(r"[^a-zA-Z0-9_-]", "-", (value or "").strip()).strip("-")
    return cleaned or fallback


def _artifact_dir(conversation_id: str | None) -> Path:
    return _EXPORT_ROOT / _safe_segment(conversation_id or "", "default")


def _safe_filename(name: str, ext: str, fallback: str) -> str:
    """Lấy base name an toàn (chống path traversal) + ép đúng đuôi định dạng."""
    base = Path((name or "").strip()).name  # bỏ mọi phần thư mục
    base = re.sub(r"[^a-zA-Z0-9_.\- ]", "", base).strip() or fallback
    # Bỏ đuôi cũ (nếu model gõ sai) rồi gắn đuôi đúng.
    stem = base[: -len(ext)] if base.lower().endswith(ext) else Path(base).stem
    stem = stem.strip() or fallback
    return f"{stem}{ext}"


def artifact_b64(conversation_id: str | None, filename: str, max_bytes: int = 6_000_000) -> str | None:
    """Đọc file đã xuất → base64 để gửi THẲNG về user qua kênh chat (không cần URL).
    Trả None nếu không có file hoặc vượt giới hạn (tránh nhồi SSE quá lớn)."""
    p = _artifact_dir(conversation_id) / Path(filename).name
    if not p.is_file() or p.stat().st_size > max_bytes:
        return None
    return base64.b64encode(p.read_bytes()).decode()


def _validate_table(columns: Any, rows: Any) -> tuple[list[str], list[list[Any]]]:
    """Kiểm tra + chuẩn hoá bảng (columns + rows) do model sinh ra. Lỗi → ValueError
    (catalog bọc thành is_error, model tự sửa)."""
    if not isinstance(columns, list) or not columns:
        raise ValueError("`columns` phải là mảng tên cột không rỗng")
    if len(columns) > _MAX_COLS:
        raise ValueError(f"quá nhiều cột ({len(columns)} > {_MAX_COLS})")
    if not isinstance(rows, list):
        raise ValueError("`rows` phải là mảng các dòng (mỗi dòng là mảng ô)")
    if len(rows) > _MAX_ROWS:
        raise ValueError(
            f"quá nhiều dòng ({len(rows)} > {_MAX_ROWS}). Hãy chia nhỏ hoặc tóm tắt dữ liệu."
        )
    cols = [str(c) for c in columns]
    norm_rows: list[list[Any]] = []
    for r in rows:
        if not isinstance(r, list):
            raise ValueError("mỗi phần tử của `rows` phải là 1 mảng các ô")
        norm_rows.append(list(r))
    return cols, norm_rows


def _meta(filename: str, size_bytes: int) -> str:
    """JSON trả về theo convention artifact chung (ChatEngine đọc để emit nút tải)."""
    return json.dumps(
        {
            "artifact": True,
            "filename": filename,
            "size_kb": round(size_bytes / 1024, 1),
            "note": (
                "File đã sẵn sàng — hệ thống ĐÃ TỰ gửi nút tải ngay dưới tin nhắn. "
                "TUYỆT ĐỐI KHÔNG in/bịa link tải (không markdown link, không /tmp/, không URL). "
                f"Chỉ cần báo tên file: {filename}."
            ),
        },
        ensure_ascii=False,
    )


class FileExportProvider:
    server_name = "file-export"
    is_mock = False  # tool THẬT — minh bạch trên trang Review

    def list_tools(self) -> list[ToolDef]:
        _table_schema = {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Tên file (không cần đuôi, hệ thống tự thêm)."},
                "columns": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Danh sách tên cột (header).",
                },
                "rows": {
                    "type": "array",
                    "items": {"type": "array", "items": {}},
                    "description": f"Các dòng dữ liệu; mỗi dòng là mảng ô theo thứ tự cột. Tối đa {_MAX_ROWS} dòng.",
                },
            },
            "required": ["columns", "rows"],
        }
        return [
            ToolDef(
                name="export_xlsx",
                description=(
                    "Xuất BẢNG DỮ LIỆU/số liệu ra file Excel (.xlsx) để user tải về. "
                    "Dùng MẶC ĐỊNH khi user muốn lưu/tải kết quả dạng bảng mà không nói rõ định dạng."
                ),
                input_schema={
                    **_table_schema,
                    "properties": {
                        **_table_schema["properties"],
                        "sheet_name": {"type": "string", "description": "Tên sheet (mặc định 'Sheet1')."},
                    },
                },
                stateful=True,  # sinh file vào artifact dir theo conversation → cần _conversation_id
            ),
            ToolDef(
                name="export_csv",
                description=(
                    "Xuất dữ liệu thô dạng CSV (UTF-8) — dùng khi user cần file để nạp sang "
                    "hệ thống/Excel khác hoặc yêu cầu rõ định dạng CSV."
                ),
                input_schema=_table_schema,
                stateful=True,  # sinh file vào artifact dir theo conversation → cần _conversation_id
            ),
            ToolDef(
                name="export_docx",
                description=(
                    "Xuất VĂN BẢN/báo cáo/tài liệu dài ra file Word (.docx) để user tải về. "
                    "Dùng khi nội dung là văn bản (báo cáo, hợp đồng, biên bản...) chứ không phải bảng."
                ),
                input_schema={
                    "type": "object",
                    "properties": {
                        "filename": {"type": "string", "description": "Tên file (không cần đuôi)."},
                        "title": {"type": "string", "description": "Tiêu đề tài liệu (tuỳ chọn)."},
                        "sections": {
                            "type": "array",
                            "description": "Các phần của tài liệu, theo thứ tự.",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "heading": {"type": "string", "description": "Tiêu đề phần (tuỳ chọn)."},
                                    "paragraphs": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                        "description": "Các đoạn văn của phần này.",
                                    },
                                },
                            },
                        },
                    },
                    "required": ["sections"],
                },
                stateful=True,  # sinh file vào artifact dir theo conversation → cần _conversation_id
            ),
        ]

    def call(self, tool_name: str, args: dict[str, Any]) -> str:
        cid = args.get("_conversation_id")
        out_dir = _artifact_dir(cid)
        out_dir.mkdir(parents=True, exist_ok=True)

        if tool_name == "export_csv":
            cols, rows = _validate_table(args.get("columns"), args.get("rows"))
            filename = _safe_filename(str(args.get("filename") or ""), ".csv", "export")
            buf = io.StringIO()
            writer = csv.writer(buf)
            writer.writerow(cols)
            writer.writerows([[("" if c is None else c) for c in r] for r in rows])
            # utf-8-sig (BOM) để Excel mở tiếng Việt không lỗi font.
            data = buf.getvalue().encode("utf-8-sig")
            (out_dir / filename).write_bytes(data)
            return _meta(filename, len(data))

        if tool_name == "export_xlsx":
            cols, rows = _validate_table(args.get("columns"), args.get("rows"))
            filename = _safe_filename(str(args.get("filename") or ""), ".xlsx", "export")
            from openpyxl import Workbook  # import lazy: chỉ nạp khi thật sự xuất xlsx

            wb = Workbook()
            ws = wb.active
            ws.title = _safe_segment(str(args.get("sheet_name") or "Sheet1"), "Sheet1")[:31]
            ws.append(cols)
            for r in rows:
                # openpyxl chỉ nhận scalar — ép ô phức tạp (list/dict) thành chuỗi.
                ws.append([c if isinstance(c, (str, int, float, bool)) or c is None else str(c) for c in r])
            target = out_dir / filename
            wb.save(target)
            return _meta(filename, target.stat().st_size)

        if tool_name == "export_docx":
            sections = args.get("sections")
            if not isinstance(sections, list) or not sections:
                raise ValueError("`sections` phải là mảng các phần không rỗng")
            if len(sections) > _MAX_ROWS:
                raise ValueError(f"quá nhiều phần ({len(sections)} > {_MAX_ROWS})")
            filename = _safe_filename(str(args.get("filename") or ""), ".docx", "tai-lieu")
            from docx import Document  # import lazy

            doc = Document()
            title = str(args.get("title") or "").strip()
            if title:
                doc.add_heading(title, level=0)
            for sec in sections:
                if not isinstance(sec, dict):
                    continue
                heading = str(sec.get("heading") or "").strip()
                if heading:
                    doc.add_heading(heading, level=1)
                for para in sec.get("paragraphs") or []:
                    doc.add_paragraph("" if para is None else str(para))
            target = out_dir / filename
            doc.save(target)
            return _meta(filename, target.stat().st_size)

        raise ValueError(f"tool không tồn tại: {tool_name}")
