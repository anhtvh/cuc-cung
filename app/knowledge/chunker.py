"""Parse tài liệu (PDF/DOCX/TXT/MD/CSV/XLSX) → text → chunk cho RAG.

Tái dùng pypdf/python-docx/openpyxl (đã là dependency). Tách riêng khỏi app/api/upload.py
để cả upload-chat lẫn ingest-knowledge dùng chung logic parse.
"""

from io import BytesIO

# Cap text trích từ bảng — file lớn vẫn được chunk_text chia nhỏ, nhưng chặn blow-up bộ nhớ.
_TABLE_MAX_CHARS = 200_000


class UnsupportedDocument(ValueError):
    """Định dạng không hỗ trợ hoặc không trích được text (vd PDF scan)."""


def extract_text(filename: str, raw: bytes) -> str:
    """Trích text thuần từ file. Raise UnsupportedDocument nếu không đọc được."""
    name = (filename or "").lower()
    if name.endswith((".txt", ".md")):
        return raw.decode("utf-8", errors="replace")
    if name.endswith(".pdf"):
        return _extract_pdf(raw)
    if name.endswith(".docx"):
        return _extract_docx(raw)
    if name.endswith(".csv"):
        return _extract_csv(raw)
    if name.endswith((".xlsx", ".xlsm")):
        return _extract_excel(raw)
    raise UnsupportedDocument(f"Định dạng không hỗ trợ cho knowledge: {filename}")


def _extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader
    try:
        reader = PdfReader(BytesIO(raw))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as e:  # noqa: BLE001
        raise UnsupportedDocument(f"Không đọc được PDF: {e}") from e
    text = "\n\n".join(p for p in pages if p.strip())
    if not text.strip():
        raise UnsupportedDocument("PDF không có text layer (ảnh scan) — cần OCR, ngoài phạm vi.")
    return text


def _extract_docx(raw: bytes) -> str:
    import docx
    try:
        doc = docx.Document(BytesIO(raw))
    except Exception as e:  # noqa: BLE001
        raise UnsupportedDocument(f"Không đọc được DOCX: {e}") from e
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not text.strip():
        raise UnsupportedDocument("DOCX rỗng — không có nội dung.")
    return text


def _extract_csv(raw: bytes) -> str:
    """CSV → text chuẩn hoá 'cell | cell' (đồng nhất với luồng chat upload)."""
    import csv
    from io import StringIO

    for enc in ("utf-8-sig", "utf-8", "cp1252"):  # utf-8-sig để nuốt BOM của Excel-export
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")

    try:
        reader = csv.reader(StringIO(text))
        lines = [" | ".join(c.strip() for c in row) for row in reader if any(c.strip() for c in row)]
        out = "\n".join(lines)
    except Exception:  # noqa: BLE001 — parse hỏng → trả raw, vẫn hữu ích hơn là fail
        out = text
    if not out.strip():
        raise UnsupportedDocument("CSV rỗng — không có dữ liệu.")
    return out[:_TABLE_MAX_CHARS]


def _extract_excel(raw: bytes) -> str:
    """Excel .xlsx/.xlsm → text. Mỗi sheet một khối, mỗi dòng 'cell | cell'."""
    try:
        from openpyxl import load_workbook
    except ImportError as e:  # pragma: no cover - openpyxl là dependency bắt buộc
        raise UnsupportedDocument("Thiếu thư viện openpyxl để đọc Excel.") from e
    try:
        wb = load_workbook(BytesIO(raw), read_only=True, data_only=True)
    except Exception as e:  # noqa: BLE001
        raise UnsupportedDocument(f"Không đọc được Excel: {e}") from e

    parts: list[str] = []
    total = 0
    for ws in wb.worksheets:
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).strip() for v in row]
            if not any(cells):
                continue
            line = " | ".join(cells).rstrip(" |")
            rows.append(line)
            total += len(line) + 1
            if total >= _TABLE_MAX_CHARS:
                break
        if rows:
            parts.append(f"# Sheet: {ws.title}\n" + "\n".join(rows))
        if total >= _TABLE_MAX_CHARS:
            parts.append("… (đã cắt bớt vì file quá lớn)")
            break
    wb.close()
    out = "\n\n".join(parts)
    if not out.strip():
        raise UnsupportedDocument("Excel rỗng — không có dữ liệu.")
    return out


def chunk_text(text: str, size: int = 1200, overlap: int = 150) -> list[str]:
    """Chia text thành các chunk ~size ký tự, ưu tiên ranh giới đoạn, overlap để giữ ngữ cảnh.

    Gộp theo đoạn (\\n\\n) cho tới khi đạt ~size; đoạn quá dài thì cắt cứng theo size.
    overlap: ký tự cuối chunk trước được nối vào đầu chunk sau (giữ liên tục ngữ nghĩa).
    """
    if not text or not text.strip():
        return []
    size = max(200, size)
    overlap = max(0, min(overlap, size // 2))

    # Tách theo đoạn; đoạn dài hơn size → cắt cứng thành nhiều mảnh.
    paragraphs: list[str] = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if len(para) <= size:
            paragraphs.append(para)
        else:
            for i in range(0, len(para), size):
                paragraphs.append(para[i:i + size])

    chunks: list[str] = []
    buf = ""
    for para in paragraphs:
        if buf and len(buf) + len(para) + 2 > size:
            chunks.append(buf)
            # nối overlap từ cuối chunk vừa chốt để giữ ngữ cảnh
            buf = (buf[-overlap:] + "\n\n" + para) if overlap else para
        else:
            buf = (buf + "\n\n" + para) if buf else para
    if buf.strip():
        chunks.append(buf)
    return chunks
